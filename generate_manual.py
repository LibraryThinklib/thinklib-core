"""
Gerador do Manual de Implementação Thinklib Core v0.2.1
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Cores por módulo ─────────────────────────────────────────────────────────
C_PLATFORMER  = RGBColor(0x1A, 0x5F, 0x8E)   # azul
C_TOPDOWN     = RGBColor(0x00, 0x89, 0x7B)   # verde-teal
C_TD          = RGBColor(0xE6, 0x51, 0x00)   # laranja
C_PAC         = RGBColor(0x6A, 0x1B, 0x9A)   # roxo
C_TELEMETRY   = RGBColor(0x37, 0x47, 0x4F)   # cinza-escuro
C_HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)   # branco
C_TIP_BG      = RGBColor(0xE3, 0xF2, 0xFD)   # azul claro (tip)
C_WARN_BG     = RGBColor(0xFF, 0xF3, 0xE0)   # laranja claro (warn)
C_BODY        = RGBColor(0x21, 0x21, 0x21)   # quase preto

def set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    hex_color = f'{color[0]:02X}{color[1]:02X}{color[2]:02X}'
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_run_color(run, color: RGBColor):
    run.font.color.rgb = color

def add_heading(doc, text, level, color):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18 if level == 1 else 13)
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    return p

def add_part_banner(doc, number, title, color):
    """Bloco PARTE N + título grande"""
    # Banner "PARTE N"
    p_banner = doc.add_paragraph()
    p_banner.paragraph_format.space_before = Pt(18)
    p_banner.paragraph_format.space_after  = Pt(2)
    banner_run = p_banner.add_run(f'  PARTE {number}  ')
    banner_run.bold = True
    banner_run.font.size = Pt(11)
    banner_run.font.color.rgb = C_HEADER_TEXT
    banner_run.font.name = 'Calibri'
    # Fundo colorido via shading no parágrafo
    pPr = p_banner._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    hex_color = f'{color[0]:02X}{color[1]:02X}{color[2]:02X}'
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)
    # Título da parte
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(4)
    p_title.paragraph_format.space_after  = Pt(4)
    t_run = p_title.add_run(title)
    t_run.bold = True
    t_run.font.size = Pt(22)
    t_run.font.color.rgb = color
    t_run.font.name = 'Calibri'
    # Linha divisória
    add_hrule(doc, color)

def add_hrule(doc, color):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    hex_color = f'{color[0]:02X}{color[1]:02X}{color[2]:02X}'
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_subheading(doc, text, color):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = C_BODY
        run.font.name = 'Calibri'
    return p

def add_properties_table(doc, header_color, rows):
    """rows: list of (prop, type_default, description)"""
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Larguras
    for i, w in enumerate([Cm(4.5), Cm(3.5), Cm(9.0)]):
        for cell in table.columns[i].cells:
            cell.width = w
    # Cabeçalho
    headers = ['Propriedade', 'Tipo / Padrão', 'Descrição']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, header_color)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = C_HEADER_TEXT
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
    # Linhas de dados
    for ri, (prop, typ, desc) in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, (val, clr) in enumerate([
            (prop, header_color),
            (typ, header_color),
            (desc, C_BODY),
        ]):
            cell = row.cells[ci]
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)
            run.font.color.rgb = clr
            run.font.name = 'Calibri'
    doc.add_paragraph()  # espaço após tabela

def add_steps(doc, steps):
    """Lista numerada de passos"""
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(step)
        run.font.size = Pt(10)
        run.font.color.rgb = C_BODY
        run.font.name = 'Calibri'
    doc.add_paragraph()

def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(item)
        run.font.size = Pt(10)
        run.font.color.rgb = C_BODY
        run.font.name = 'Calibri'
    doc.add_paragraph()

def add_tip(doc, text, bg=None, icon='💡'):
    if bg is None:
        bg = C_TIP_BG
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    # Remove bordas
    for cell in table.rows[0].cells:
        set_cell_bg(cell, bg)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'none')
            tcBorders.append(b)
        tcPr.append(tcBorders)
    icon_cell = table.rows[0].cells[0]
    icon_cell.width = Cm(0.8)
    p = icon_cell.paragraphs[0]
    r = p.add_run(icon)
    r.font.size = Pt(14)
    text_cell = table.rows[0].cells[1]
    p2 = text_cell.paragraphs[0]
    run = p2.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = C_BODY
    run.font.name = 'Calibri'
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_after = Pt(4)

def add_page_break(doc):
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# GERAR DOCUMENTO
# ═══════════════════════════════════════════════════════════════════════════════
doc = Document()

# Margens
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ─── CAPA ─────────────────────────────────────────────────────────────────────
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(80)
r = p_title.add_run('THINKLIB CORE')
r.bold = True
r.font.size = Pt(36)
r.font.color.rgb = C_PLATFORMER
r.font.name = 'Calibri'

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_before = Pt(6)
r2 = p_sub.add_run('Manual de Implementação de Mecânicas')
r2.font.size = Pt(18)
r2.font.color.rgb = C_PLATFORMER
r2.font.name = 'Calibri'

p_ver = doc.add_paragraph()
p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_ver.paragraph_format.space_before = Pt(10)
r3 = p_ver.add_run('Versão 0.2.1  •  Unity 2022.3+  •  C#')
r3.font.size = Pt(11)
r3.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
r3.font.name = 'Calibri'

add_hrule(doc, C_PLATFORMER)

p_cats = doc.add_paragraph()
p_cats.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_cats.paragraph_format.space_before = Pt(10)
r4 = p_cats.add_run('Platformer  •  Topdown  •  Tower Defense  •  Point & Click')
r4.font.size = Pt(12)
r4.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
r4.font.name = 'Calibri'

add_page_break(doc)

# ─── INTRODUÇÃO ───────────────────────────────────────────────────────────────
add_heading(doc, 'Introdução', 1, C_PLATFORMER)
add_body(doc,
    'O Thinklib Core é um pacote Unity (UPM) de mecânicas reutilizáveis e modulares para '
    'aceleração do desenvolvimento de jogos. Este manual descreve como instalar o pacote e '
    'integrar cada mecânica em um projeto Unity 2022.3 ou superior.')

add_subheading(doc, 'Instalação via UPM', C_PLATFORMER)
add_body(doc,
    'No Unity, acesse Window → Package Manager, clique em + e escolha Add package from Git URL. '
    'Cole o endereço abaixo e confirme:')
p_code = doc.add_paragraph('    https://github.com/LibraryThinklib/thinklib-core.git#v0.2.1')
p_code.paragraph_format.left_indent = Cm(1)
for run in p_code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = C_PLATFORMER
add_body(doc, 'O Unity fará o download automático das dependências TextMesh Pro (3.0.9) e UGUI (1.0.0).')

add_subheading(doc, 'Importando Prefabs', C_PLATFORMER)
add_body(doc,
    'Após a instalação, acesse o menu Thinklib → Import Resources para copiar os prefabs prontos '
    'para Assets/Thinklib/Resources/Prefabs/ no seu projeto. Esse passo é opcional: os scripts '
    'funcionam independentemente dos prefabs, mas os prefabs já trazem referências de sprites, '
    'materiais e animações configuradas.')

add_subheading(doc, 'Estrutura do Pacote', C_PLATFORMER)
add_properties_table(doc, C_PLATFORMER, [
    ('Runtime/',      'Pasta',   'Scripts de gameplay: Platformer, Topdown, TowerDefense, Common, Point & Click'),
    ('Editor/',       'Pasta',   'Scripts de editor: inspectors personalizados e menus Thinklib →'),
    ('package.json',  'Arquivo', 'Metadados UPM (nome: com.thinklib.core)'),
    ('CHANGELOG.md',  'Arquivo', 'Histórico completo de versões'),
])
add_tip(doc, 'Todos os componentes aparecem no menu Component → Thinklib → … do Unity, facilitando a adição via Inspector.')
add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# PARTE 1 — PLATFORMER
# ═══════════════════════════════════════════════════════════════════════════════
add_part_banner(doc, 1, 'Platformer', C_PLATFORMER)
add_body(doc,
    'O módulo Platformer cobre todos os sistemas necessários para um jogo de plataforma 2D: '
    'movimentação, pulo, combate corpo a corpo e à distância, inimigos com inteligência básica, '
    'ambiente interativo e sistema de vida.')

# 1.1 Movimentação
add_heading(doc, '1.1  Movimentação — PlatformerMovementController', 2, C_PLATFORMER)
add_body(doc,
    'O PlatformerMovementController gerencia o deslocamento horizontal do personagem no eixo X '
    'usando o Rigidbody2D do Unity. Ele lê o eixo Horizontal do Input System legado e aplica a '
    'velocidade diretamente no componente físico.')
add_subheading(doc, 'Componentes Necessários no GameObject', C_PLATFORMER)
add_bullets(doc, [
    'Rigidbody2D — com Gravity Scale entre 2 e 4 para uma física satisfatória de plataforma',
    'Collider2D (BoxCollider2D ou CapsuleCollider2D) — ajustado ao sprite do personagem',
    'PlatformerMovementController — adicionado via Component → Thinklib → Platformer → Movement',
])
add_subheading(doc, 'Propriedades do Inspector', C_PLATFORMER)
add_properties_table(doc, C_PLATFORMER, [
    ('Move Speed',  'float | 5',        'Velocidade de deslocamento horizontal em unidades/s'),
    ('Flip Sprite', 'bool | true',      'Se verdadeiro, espelha o sprite conforme a direção de movimento'),
    ('Animator',    'Animator | null',  'Referência opcional ao Animator; define o parâmetro "Speed" automaticamente'),
])
add_subheading(doc, 'Passo a Passo', C_PLATFORMER)
add_steps(doc, [
    'Crie um GameObject vazio chamado Player.',
    'Adicione Rigidbody2D. Defina Gravity Scale = 3 e Constraints → Freeze Rotation Z = true.',
    'Adicione um CapsuleCollider2D e ajuste ao tamanho do sprite.',
    'Adicione PlatformerMovementController via menu Component → Thinklib → Platformer → Movement.',
    'Configure Move Speed no Inspector (sugestão: 5 para plataforma padrão, 7 para jogos mais rápidos).',
    'Se tiver Animator, arraste-o para o campo Animator e crie um parâmetro float chamado Speed.',
])
add_tip(doc, 'O script utiliza Input.GetAxis("Horizontal"), compatível com o Input System clássico do Unity. '
             'Para usar o novo Input System, substitua a leitura pelo InputAction correspondente.')

# 1.2 Pulo
add_heading(doc, '1.2  Pulo — PlatformerJumpController', 2, C_PLATFORMER)
add_body(doc,
    'O PlatformerJumpController detecta se o personagem está no chão e aplica uma força vertical '
    'ao Rigidbody2D quando o botão de pulo é pressionado. O controle de "está no chão" é feito '
    'por um raycast ou overlap circle configurável.')
add_subheading(doc, 'Componentes Necessários no GameObject', C_PLATFORMER)
add_bullets(doc, [
    'Rigidbody2D — o mesmo do PlatformerMovementController',
    'PlatformerJumpController — no mesmo GameObject do controlador de movimento',
])
add_subheading(doc, 'Propriedades do Inspector', C_PLATFORMER)
add_properties_table(doc, C_PLATFORMER, [
    ('Jump Force',    'float | 10',       'Força de impulso vertical aplicada ao Rigidbody2D'),
    ('Ground Layer',  'LayerMask',        'Máscara de layers consideradas "chão" para o raycast'),
    ('Ground Check',  'Transform | null', 'Ponto de origem do raycast (geralmente um filho vazio na base do collider)'),
    ('Ground Radius', 'float | 0.2',      'Raio do overlap circle de checagem de chão'),
    ('Animator',      'Animator | null',  'Referência ao Animator; define o parâmetro bool "IsGrounded"'),
])
add_subheading(doc, 'Passo a Passo', C_PLATFORMER)
add_steps(doc, [
    'No GameObject Player, crie um filho vazio chamado GroundCheck e posicione-o na base do collider (y ligeiramente negativo).',
    'Adicione PlatformerJumpController ao Player.',
    'No campo Ground Check, arraste o objeto GroundCheck.',
    'No campo Ground Layer, selecione a layer do seu tilemap ou plataformas (ex.: "Ground").',
    'Ajuste Jump Force — valores entre 8 e 15 são comuns para jogos de plataforma 2D.',
])
add_tip(doc, 'Certifique-se de que os objetos de chão estão na layer correta. Sem isso, o personagem poderá pular infinitamente no ar.', C_WARN_BG, '⚠')

# 1.3 Sistema de Vida
add_heading(doc, '1.3  Sistema de Vida — LifeUIBar e LifeUIIcons', 2, C_PLATFORMER)
add_body(doc, 'O Thinklib oferece dois estilos de exibição de vida: uma barra de progresso (LifeUIBar) e ícones individuais (LifeUIIcons), ambos em Runtime/Common/LifeSystem.')
add_subheading(doc, 'LifeUIBar — Barra de Vida', C_PLATFORMER)
add_properties_table(doc, C_TOPDOWN, [
    ('Max Health',   'int | 100',      'Valor máximo de vida do personagem'),
    ('Current HP',   'int | (Max)',    'Vida atual; altere via código: lifeBar.TakeDamage(amount)'),
    ('Fill Image',   'Image | null',   'Componente Image do Unity UI usado como preenchimento da barra'),
    ('Smooth Speed', 'float | 5',      'Velocidade da animação de lerp da barra'),
])
add_subheading(doc, 'LifeUIIcons — Ícones de Vida (estilo corações)', C_PLATFORMER)
add_properties_table(doc, C_TOPDOWN, [
    ('Max Health',      'int | 3',          'Quantidade máxima de ícones (corações)'),
    ('Icon Full',       'Sprite | null',    'Sprite exibido para vida cheia'),
    ('Icon Empty',      'Sprite | null',    'Sprite exibido para vida vazia'),
    ('Icon Container',  'Transform | null', 'Pai dos ícones gerados dinamicamente'),
])
add_subheading(doc, 'Passo a Passo — LifeUIBar', C_PLATFORMER)
add_steps(doc, [
    'Na cena, crie um Canvas (modo Screen Space – Overlay).',
    'Dentro do Canvas, crie um GameObject com um Image (background da barra) e um filho Image (fill).',
    'Adicione LifeUIBar ao GameObject da barra.',
    'Arraste a Image filho (fill) para o campo Fill Image.',
    'No script do personagem, obtenha a referência e chame TakeDamage(10) para reduzir a vida.',
])
add_subheading(doc, 'Passo a Passo — LifeUIIcons', C_PLATFORMER)
add_steps(doc, [
    'Crie um Canvas com um HorizontalLayoutGroup como container dos ícones.',
    'Adicione LifeUIIcons ao container.',
    'Atribua os sprites Icon Full e Icon Empty.',
    'Defina Max Health com o número de corações desejado.',
])

# 1.4 Efeitos
add_heading(doc, '1.4  Efeitos — DeathEffect e PlayerHurtEffect', 2, C_PLATFORMER)
add_body(doc, 'Esses dois scripts em Runtime/Common/Effects disparam animações ou partículas nos eventos de dano e morte do personagem.')
add_subheading(doc, 'DeathEffect', C_PLATFORMER)
add_properties_table(doc, C_PLATFORMER, [
    ('Death Anim Trigger', 'string | "Die"',  'Nome do trigger no Animator acionado na morte'),
    ('Destroy Delay',      'float | 1.0',     'Tempo em segundos antes de destruir o GameObject'),
    ('Death Sound',        'AudioClip | null','Clipe de áudio tocado ao morrer'),
])
add_subheading(doc, 'PlayerHurtEffect', C_PLATFORMER)
add_properties_table(doc, C_PLATFORMER, [
    ('Hurt Anim Trigger',  'string | "Hurt"',   'Nome do trigger no Animator acionado ao tomar dano'),
    ('Invincibility Time', 'float | 0.5',        'Duração da invencibilidade temporária após dano (em segundos)'),
    ('Flash Color',        'Color | vermelho',   'Cor do efeito de flash no SpriteRenderer'),
    ('Flash Count',        'int | 3',            'Número de flashes durante a invencibilidade'),
])
add_subheading(doc, 'Passo a Passo — Integração com o Sistema de Vida', C_PLATFORMER)
add_steps(doc, [
    'Adicione DeathEffect e PlayerHurtEffect ao mesmo GameObject do personagem.',
    'No script de vida, obtenha as referências via GetComponent.',
    'Chame GetComponent<PlayerHurtEffect>().PlayHurt() ao tomar dano.',
    'Chame GetComponent<DeathEffect>().PlayDeath() ao morrer.',
])

# 1.5 Combate
add_heading(doc, '1.5  Combate — Melee e Shooter', 2, C_PLATFORMER)
add_body(doc, 'O módulo de combate do Platformer oferece três scripts principais: ataque corpo a corpo, ataque à distância e controle de projéteis.')
add_subheading(doc, 'PlayerMeleeAttackController', C_PLATFORMER)
add_properties_table(doc, C_PLATFORMER, [
    ('Attack Damage',  'int | 10',          'Dano causado por ataque melee'),
    ('Attack Range',   'float | 1.0',       'Alcance do hitbox de melee em unidades Unity'),
    ('Attack Cooldown','float | 0.5',       'Intervalo mínimo entre ataques (segundos)'),
    ('Attack Layer',   'LayerMask',         'Layers que recebem dano (ex.: Enemy)'),
    ('Anim Trigger',   'string | "Attack"', 'Nome do trigger no Animator'),
])
add_subheading(doc, 'PlayerShooterController', C_PLATFORMER)
add_properties_table(doc, C_PLATFORMER, [
    ('Projectile Prefab', 'GameObject | null', 'Prefab do projétil disparado'),
    ('Fire Point',        'Transform | null',  'Ponto de origem dos projéteis (filho do personagem)'),
    ('Fire Rate',         'float | 0.3',       'Intervalo entre disparos em segundos'),
    ('Projectile Speed',  'float | 10',        'Velocidade dos projéteis em unidades/s'),
    ('Fire Sound',        'AudioClip | null',  'Clipe de áudio do disparo'),
])
add_subheading(doc, 'ProjectileDamageDealer', C_PLATFORMER)
add_properties_table(doc, C_PLATFORMER, [
    ('Damage',         'int | 10',           'Dano aplicado ao alvo ao colidir'),
    ('Target Layer',   'LayerMask',          'Layers que recebem dano'),
    ('Destroy On Hit', 'bool | true',        'Destrói o projétil ao colidir'),
    ('Hit Effect',     'GameObject | null',  'Prefab de partícula/efeito instanciado ao acertar'),
])
add_subheading(doc, 'Passo a Passo — Ataque à Distância', C_PLATFORMER)
add_steps(doc, [
    'Crie um prefab de projétil com SpriteRenderer, Rigidbody2D (kinematic) e Collider2D (Is Trigger = true).',
    'Adicione ProjectileDamageDealer ao prefab e configure Damage e Target Layer.',
    'No Player, adicione PlayerShooterController.',
    'Crie um filho vazio chamado FirePoint na ponta da arma/mão do personagem.',
    'Arraste o prefab do projétil para Projectile Prefab e o FirePoint para Fire Point.',
    'Pressione o botão de disparo configurado (padrão: Fire1 / Ctrl esquerdo) para testar.',
])

# 1.6 Inimigos
add_heading(doc, '1.6  Inimigos — PatrollerAI e EnemyShooterAI', 2, C_PLATFORMER)
add_body(doc, 'O Thinklib oferece dois arquétipos de inimigos para Platformer: o Patroller (patrulheiro que causa dano por contato) e o Shooter (atirador com visão de linha).')
add_subheading(doc, 'PatrollerAI + DamageOnTouch', C_PLATFORMER)
add_properties_table(doc, C_PLATFORMER, [
    ('Patrol Points', 'Transform[]',    'Array de pontos que o inimigo percorre em loop'),
    ('Move Speed',    'float | 2',      'Velocidade de patrulha'),
    ('Wait Time',     'float | 1',      'Tempo de espera em cada ponto de patrulha'),
    ('Flip On Turn',  'bool | true',    'Espelha o sprite ao inverter direção'),
    ('Damage',        'int | 10',       'Dano causado ao jogador por colisão (DamageOnTouch)'),
    ('Player Layer',  'LayerMask',      'Layer do jogador para detecção de contato'),
    ('Knockback Force','Vector2 | (3,5)','Força de knockback aplicada ao jogador'),
])
add_subheading(doc, 'EnemyShooterAI', C_PLATFORMER)
add_properties_table(doc, C_PLATFORMER, [
    ('Projectile Prefab', 'GameObject | null', 'Prefab do projétil do inimigo'),
    ('Fire Point',        'Transform | null',  'Ponto de origem dos projéteis'),
    ('Detection Range',   'float | 8',         'Alcance de detecção do jogador em unidades'),
    ('Fire Rate',         'float | 2',         'Intervalo entre disparos do inimigo'),
    ('Player Layer',      'LayerMask',         'Layer do jogador para detecção visual'),
])
add_subheading(doc, 'Passo a Passo — Configurando Pontos de Patrulha', C_PLATFORMER)
add_steps(doc, [
    'Crie um GameObject pai chamado PatrolRoute.',
    'Adicione filhos vazios (P1, P2, P3…) nas posições desejadas do percurso.',
    'No Inspector do PatrollerAI, defina o tamanho do array Patrol Points e arraste cada ponto.',
])
add_tip(doc, 'Use Gizmos: o PatrollerAI desenha a rota de patrulha como linhas coloridas no Scene View quando selecionado.')

# 1.7 Ambiente — Plataformas
add_heading(doc, '1.7  Ambiente — MovingPlatform e TimedPlatform', 2, C_PLATFORMER)
add_subheading(doc, 'MovingPlatform', C_PLATFORMER)
add_body(doc, 'Plataforma que se move continuamente entre dois ou mais pontos. O personagem é automaticamente filho da plataforma durante o contato, evitando deslizamento.')
add_properties_table(doc, C_PLATFORMER, [
    ('Waypoints', 'Transform[]',  'Pontos pelos quais a plataforma se move em loop'),
    ('Speed',     'float | 2',    'Velocidade de translação em unidades/s'),
    ('Wait Time', 'float | 0',    'Pausa em cada waypoint (0 = sem pausa)'),
    ('Loop',      'bool | true',  'Se verdadeiro, volta ao primeiro waypoint ao chegar no último'),
])
add_subheading(doc, 'TimedPlatform', C_PLATFORMER)
add_body(doc, 'Plataforma que desaparece e reaparece em intervalos configuráveis.')
add_properties_table(doc, C_PLATFORMER, [
    ('Visible Time',  'float | 3',    'Tempo em segundos que a plataforma fica visível/sólida'),
    ('Hidden Time',   'float | 2',    'Tempo em segundos que a plataforma fica invisível/passável'),
    ('Start Visible', 'bool | true',  'Estado inicial da plataforma ao iniciar a cena'),
    ('Warning Time',  'float | 0.5',  'Tempo antes de sumir em que o sprite pisca como aviso'),
])
add_subheading(doc, 'Passo a Passo', C_PLATFORMER)
add_steps(doc, [
    'Crie um GameObject para a plataforma com SpriteRenderer e Collider2D.',
    'Adicione MovingPlatform ou TimedPlatform via Component → Thinklib → Platformer → Environment.',
    'Para MovingPlatform, crie filhos vazios (W1, W2…) como waypoints e atribua ao array Waypoints.',
    'Para TimedPlatform, ajuste Visible Time e Hidden Time conforme a dificuldade desejada.',
])
add_tip(doc, 'Para plataformas que o jogador pode atravessar de baixo (one-way), use o componente PlatformEffector2D do Unity junto com o Collider2D da plataforma.')

# 1.8 Ambiente — RewardChest e SawHazard (NOVO)
add_heading(doc, '1.8  Ambiente — RewardChest e SawHazard', 2, C_PLATFORMER)
add_body(doc, 'Objetos interativos de ambiente para enriquecer fases de plataforma. Adicionados na v0.2.1.')
add_subheading(doc, 'RewardChest — Baú de Recompensa', C_PLATFORMER)
add_body(doc, 'Ao toque do jogador (tag Player), cura vida via LifeSystemController.Heal(). Troca sprite ao ser usado e desativa o collider para não ser ativado novamente.')
add_properties_table(doc, C_PLATFORMER, [
    ('Health To Give', 'int | 1',         'Quantidade de vida curada ao coletar o baú'),
    ('Open Sprite',    'Sprite | null',   'Sprite exibido após o baú ser aberto (opcional)'),
])
add_subheading(doc, 'SawHazard — Perigo de Serra', C_PLATFORMER)
add_body(doc, 'Causa dano ao jogador via LifeSystemController.TakeDamage(). Suporta modo de patrulha entre dois pontos com Gizmos de debug no Scene View.')
add_properties_table(doc, C_PLATFORMER, [
    ('Damage Amount', 'int | 1',       'Dano causado ao tocar o jogador'),
    ('Should Move',   'bool | false',  'Se verdadeiro, a serra se move entre Point A e Point B'),
    ('Point A',       'Vector2',       'Posição inicial da patrulha'),
    ('Point B',       'Vector2',       'Posição final da patrulha'),
    ('Move Speed',    'float | 3.0',   'Velocidade de deslocamento da serra'),
])
add_subheading(doc, 'Passo a Passo', C_PLATFORMER)
add_steps(doc, [
    'Crie um GameObject com SpriteRenderer e Collider2D (Is Trigger = true).',
    'Adicione RewardChest ou SawHazard via Component → Thinklib → Game.',
    'Garanta que o personagem do jogador possui o script LifeSystemController e a tag "Player".',
    'Para SawHazard com patrulha: marque Should Move = true e defina Point A e Point B no Inspector.',
    'Os Gizmos da serra aparecem no Scene View ao selecionar o objeto, facilitando o ajuste do percurso.',
])

# 1.9 Coletáveis
add_heading(doc, '1.9  Coletáveis — CollectibleItem e GameManager', 2, C_PLATFORMER)
add_subheading(doc, 'CollectibleItem', C_PLATFORMER)
add_body(doc, 'Script que detecta colisão com o jogador e notifica o GameManager ao ser coletado.')
add_properties_table(doc, C_PLATFORMER, [
    ('Item Type',      'string | "coin"',     'Identificador do tipo do item (coin, key, star, etc.)'),
    ('Item Value',     'int | 1',             'Valor numérico adicionado ao contador no GameManager'),
    ('Collect Sound',  'AudioClip | null',    'Áudio tocado ao coletar'),
    ('Collect Effect', 'GameObject | null',   'Prefab de partícula instanciado ao coletar'),
    ('Player Layer',   'LayerMask',           'Layer do jogador para detecção de trigger'),
])
add_subheading(doc, 'GameManager (Platformer)', C_PLATFORMER)
add_body(doc, 'Singleton que controla o estado global do nível: pontuação, coletáveis e condição de vitória.')
add_properties_table(doc, C_PLATFORMER, [
    ('Score UI',      'TMP_Text | null',       'Texto TextMeshPro que exibe a pontuação'),
    ('Win Condition', 'int | 10',              'Número de coletáveis necessários para vencer o nível'),
    ('Win Screen',    'GameObject | null',     'Painel de UI exibido ao atingir a condição de vitória'),
])
add_subheading(doc, 'Passo a Passo', C_PLATFORMER)
add_steps(doc, [
    'Crie um GameObject GameManager na cena e adicione o script GameManager.',
    'Configure os campos Score UI e Win Screen.',
    'Para cada moeda/item, crie um prefab com Collider2D (Is Trigger = true) e CollectibleItem.',
    'Defina Item Type e Item Value conforme a categoria do item.',
])

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# PARTE 2 — TOPDOWN
# ═══════════════════════════════════════════════════════════════════════════════
add_part_banner(doc, 2, 'Topdown', C_TOPDOWN)
add_body(doc,
    'O módulo Topdown cobre jogos com visão de cima (câmera ortográfica 2D): movimentação em '
    '4 ou 8 direções com Blend Trees, combate melee e shooter, inimigos com IA e sistema de NPC '
    'com diálogo.')

# 2.1 Movimentação
add_heading(doc, '2.1  Movimentação — TopdownMovementController', 2, C_TOPDOWN)
add_body(doc, 'O TopdownMovementController gerencia o deslocamento nos eixos X e Y usando Rigidbody2D. Ele alimenta os parâmetros de Blend Tree do Animator para animações de 8 direções.')
add_subheading(doc, 'Componentes Necessários no GameObject', C_TOPDOWN)
add_bullets(doc, [
    'Rigidbody2D — Gravity Scale = 0 (visão de cima não usa gravidade), Collision Detection = Continuous',
    'Collider2D — BoxCollider2D ou CircleCollider2D ajustado ao sprite',
    'TopdownMovementController — Component → Thinklib → Topdown → Movement',
])
add_subheading(doc, 'Propriedades do Inspector', C_TOPDOWN)
add_properties_table(doc, C_TOPDOWN, [
    ('Move Speed',      'float | 5',          'Velocidade de deslocamento em todas as direções'),
    ('Animator',        'Animator | null',    'Referência ao Animator para Blend Trees'),
    ('Blend Param X',   'string | "MoveX"',  'Nome do parâmetro float no Animator para o eixo X'),
    ('Blend Param Y',   'string | "MoveY"',  'Nome do parâmetro float no Animator para o eixo Y'),
    ('Normalize Input', 'bool | true',        'Normaliza o vetor de entrada para evitar velocidade diagonal maior'),
])
add_subheading(doc, 'Passo a Passo — Configurando a Blend Tree no Animator', C_TOPDOWN)
add_steps(doc, [
    'No Animator Controller, crie dois parâmetros float: MoveX e MoveY.',
    'Crie um estado com uma 2D Simple Directional Blend Tree.',
    'Configure os parâmetros X e Y como MoveX e MoveY.',
    'Adicione os clips de animação: Walk_Up, Walk_Down, Walk_Left, Walk_Right (e diagonais se houver).',
    'No Inspector do TopdownMovementController, preencha os campos Blend Param X e Y.',
])
add_tip(doc, 'Para jogos Topdown com apenas 4 direções, use uma 1D Blend Tree com um único parâmetro MoveX e trate as animações de cima/baixo separadamente.')

# 2.2 Combate Topdown
add_heading(doc, '2.2  Combate Topdown — Melee e Shooter', 2, C_TOPDOWN)
add_subheading(doc, 'PlayerTopdownMeleeAttackController', C_TOPDOWN)
add_properties_table(doc, C_TOPDOWN, [
    ('Attack Damage',  'int | 15',          'Dano do ataque melee'),
    ('Attack Range',   'float | 1.5',       'Alcance do hitbox em unidades'),
    ('Attack Offset',  'Vector2 | (1, 0)',  'Deslocamento do hitbox relativo ao personagem'),
    ('Attack Cooldown','float | 0.4',       'Tempo de recarga entre ataques'),
    ('Attack Layer',   'LayerMask',         'Layers que recebem dano'),
    ('Anim Trigger',   'string | "Attack"', 'Trigger no Animator'),
])
add_subheading(doc, 'PlayerTopdownShooterController', C_TOPDOWN)
add_properties_table(doc, C_TOPDOWN, [
    ('Projectile Prefab', 'GameObject | null', 'Prefab do projétil'),
    ('Fire Point',        'Transform | null',  'Ponto de origem (filho do personagem)'),
    ('Fire Rate',         'float | 0.5',       'Intervalo entre disparos'),
    ('Projectile Speed',  'float | 12',        'Velocidade do projétil'),
    ('Aim Mode',          'enum | Mouse',      'Mouse: mira pelo cursor. Direction: mira pela direção de movimento'),
])
add_subheading(doc, 'Passo a Passo — Configurando Aim Mode = Mouse', C_TOPDOWN)
add_steps(doc, [
    'Certifique-se de que a câmera tem o tag "MainCamera".',
    'O script converte automaticamente a posição do cursor para coordenadas de mundo.',
    'O projétil é disparado na direção do cursor ao pressionar Fire1.',
    'Para Aim Mode = Direction, o projétil segue a última direção de movimento do personagem.',
])

# 2.3 Inimigos Topdown
add_heading(doc, '2.3  Inimigos Topdown — Patroller e Shooter', 2, C_TOPDOWN)
add_subheading(doc, 'TopdownPatrollerAI + TopdownDamageOnTouch', C_TOPDOWN)
add_properties_table(doc, C_TOPDOWN, [
    ('Patrol Points', 'Transform[]',   'Pontos de patrulha no plano XY'),
    ('Move Speed',    'float | 2',     'Velocidade de patrulha'),
    ('Wait Time',     'float | 1',     'Pausa em cada waypoint'),
    ('Chase Player',  'bool | false',  'Se verdadeiro, persegue o jogador ao detectá-lo'),
    ('Chase Range',   'float | 5',     'Alcance de detecção para início da perseguição'),
])
add_subheading(doc, 'TopdownEnemyShooterAI', C_TOPDOWN)
add_properties_table(doc, C_TOPDOWN, [
    ('Projectile Prefab', 'GameObject | null', 'Prefab do projétil'),
    ('Fire Point',        'Transform | null',  'Origem dos projéteis'),
    ('Detection Range',   'float | 6',         'Alcance de detecção visual do jogador'),
    ('Fire Rate',         'float | 2.5',       'Intervalo entre disparos'),
    ('Rotate To Player',  'bool | true',       'Rotaciona o GameObject em direção ao jogador'),
])
add_subheading(doc, 'Passo a Passo — Inimigo com Perseguição + Tiro', C_TOPDOWN)
add_steps(doc, [
    'Adicione TopdownPatrollerAI com Chase Player = true.',
    'Adicione TopdownEnemyShooterAI ao mesmo GameObject.',
    'O patroller vai perseguir até o alcance de disparo, então o shooter assumirá o ataque.',
    'Certifique-se de que Chase Range do Patroller é maior que Detection Range do Shooter.',
])

# 2.4 NPC
add_heading(doc, '2.4  NPC com Diálogo — TopdownNPCController e DialogueBubble', 2, C_TOPDOWN)
add_body(doc, 'O sistema de NPC exibe um balão de fala acima do personagem ao interagir com o jogador, suportando múltiplas linhas de texto em sequência.')
add_subheading(doc, 'TopdownNPCController', C_TOPDOWN)
add_properties_table(doc, C_TOPDOWN, [
    ('Dialogue Lines',  'string[]',              'Array de falas do NPC em ordem'),
    ('Interact Key',    'KeyCode | E',           'Tecla de interação do jogador'),
    ('Interact Range',  'float | 2',             'Distância máxima de interação'),
    ('Dialogue Bubble', 'DialogueBubble | null', 'Referência ao componente de balão'),
    ('Player Layer',    'LayerMask',             'Layer do jogador para detecção de proximidade'),
])
add_subheading(doc, 'DialogueBubble', C_TOPDOWN)
add_properties_table(doc, C_TOPDOWN, [
    ('Bubble Root',   'GameObject | null', 'Root do balão de UI (ativado/desativado)'),
    ('Text Field',    'TMP_Text | null',   'Componente TextMeshPro onde o texto é exibido'),
    ('Char Delay',    'float | 0.04',      'Delay entre cada caractere no efeito typewriter'),
    ('Bubble Offset', 'Vector3 | (0,1.5,0)','Posição do balão relativa ao NPC no World Space'),
])
add_subheading(doc, 'Passo a Passo', C_TOPDOWN)
add_steps(doc, [
    'Crie o GameObject NPC com sprite e adicione TopdownNPCController.',
    'Crie um filho Canvas (World Space) com um Image de fundo e um TMP_Text.',
    'Adicione DialogueBubble ao Canvas filho e configure Bubble Root e Text Field.',
    'Arraste o DialogueBubble para o campo do TopdownNPCController.',
    'Preencha o array Dialogue Lines com as falas do NPC.',
    'Ao pressionar E perto do NPC, as falas avançam uma a uma e o balão fecha ao terminar.',
])

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# PARTE 3 — TOWER DEFENSE
# ═══════════════════════════════════════════════════════════════════════════════
add_part_banner(doc, 3, 'Tower Defense', C_TD)
add_body(doc,
    'O módulo Tower Defense oferece um ciclo completo: spawning de inimigos, progressão por '
    'caminho, sistema de torres com compra, colocação e upgrade, gestão de recursos e condição '
    'de derrota.')

# 3.1 Inimigos TD
add_heading(doc, '3.1  Inimigos — EnemySpawner, EnemyPath e EnemyHealth', 2, C_TD)
add_subheading(doc, 'EnemySpawner', C_TD)
add_body(doc, 'Gerencia as ondas de inimigos, instanciando-os em intervalos configuráveis.')
add_properties_table(doc, C_TD, [
    ('Enemy Prefab',     'GameObject | null', 'Prefab do inimigo a ser spawnado'),
    ('Spawn Point',      'Transform | null',  'Posição inicial dos inimigos (início do caminho)'),
    ('Enemies Per Wave', 'int | 5',           'Quantidade de inimigos por onda'),
    ('Spawn Interval',   'float | 1.5',       'Intervalo em segundos entre spawns de uma mesma onda'),
    ('Wave Interval',    'float | 5',         'Pausa em segundos entre ondas'),
    ('Total Waves',      'int | 3',           'Número total de ondas'),
    ('Auto Start',       'bool | true',       'Inicia a primeira onda automaticamente ao carregar a cena'),
])
add_subheading(doc, 'EnemyPath', C_TD)
add_properties_table(doc, C_TD, [
    ('Waypoints',   'Transform[]', 'Sequência de pontos do caminho (início ao fim)'),
    ('Move Speed',  'float | 3',   'Velocidade de deslocamento pelo caminho'),
    ('Reached End', 'UnityEvent',  'Evento disparado quando o inimigo chega ao objetivo'),
])
add_subheading(doc, 'EnemyHealth', C_TD)
add_properties_table(doc, C_TD, [
    ('Max Health',   'int | 50',           'Vida máxima do inimigo'),
    ('Reward',       'int | 10',           'Recursos concedidos ao jogador ao matar o inimigo'),
    ('Death Effect', 'GameObject | null',  'Prefab de efeito instanciado ao morrer'),
    ('Health Bar',   'Image | null',       'Barra de vida sobre o inimigo (opcional)'),
])
add_subheading(doc, 'Passo a Passo — Configurando o Caminho', C_TD)
add_steps(doc, [
    'Crie um GameObject vazio Path com filhos Waypoint_01, Waypoint_02 etc. posicionados no mapa.',
    'No EnemyPath do prefab do inimigo, popule o array Waypoints com os filhos em ordem.',
    'Defina a posição do último waypoint sobre o objetivo que os inimigos devem alcançar.',
    'Configure o EnemySpawner com o prefab do inimigo e o Spawn Point (início do caminho).',
])
add_tip(doc, 'O array Waypoints do EnemyPath deve ser configurado no prefab, não na cena, para que todos os inimigos spawnados já conheçam o caminho.')

# 3.2 Torres
add_heading(doc, '3.2  Torres — TowerPlacement, TowerShooter e TowerUpgrade', 2, C_TD)
add_subheading(doc, 'TowerPlacement', C_TD)
add_properties_table(doc, C_TD, [
    ('Tower Prefab',  'GameObject | null',    'Prefab da torre a ser posicionada'),
    ('Tower Cost',    'int | 50',             'Custo em recursos para comprar a torre'),
    ('Valid Layer',   'LayerMask',            'Layer dos tiles/células válidos para posicionamento'),
    ('Invalid Color', 'Color | vermelho',     'Cor do preview quando a posição é inválida'),
    ('Valid Color',   'Color | verde',        'Cor do preview quando a posição é válida'),
    ('Player Score',  'PlayerScore | null',   'Referência ao componente de recursos do jogador'),
])
add_subheading(doc, 'TowerShooter', C_TD)
add_properties_table(doc, C_TD, [
    ('Range',            'float | 5',           'Raio de detecção de inimigos'),
    ('Fire Rate',        'float | 1',           'Disparos por segundo'),
    ('Projectile',       'GameObject | null',   'Prefab do projétil da torre'),
    ('Fire Point',       'Transform | null',    'Ponto de origem dos projéteis'),
    ('Target Mode',      'enum | First',        'First: inimigo mais à frente. Nearest: mais próximo. Strongest: mais vida'),
    ('Rotate To Target', 'bool | true',         'A torre rotaciona em direção ao alvo'),
])
add_subheading(doc, 'TowerUpgrade', C_TD)
add_properties_table(doc, C_TD, [
    ('Upgrade Levels', 'UpgradeData[]',        'Array com dados de cada nível (custo, dano, alcance, cadência)'),
    ('Current Level',  'int | 0',             'Nível atual da torre (somente leitura no Inspector)'),
    ('Upgrade UI',     'GameObject | null',   'Painel de UI com botão de upgrade'),
    ('Player Score',   'PlayerScore | null',  'Referência ao componente de recursos'),
])
add_subheading(doc, 'Passo a Passo — Torre Básica', C_TD)
add_steps(doc, [
    'Crie um prefab Tower com SpriteRenderer, Collider2D e adicione TowerShooter.',
    'Configure Range, Fire Rate e Projectile no Inspector.',
    'Na cena, adicione um GameObject TowerShop com TowerPlacement.',
    'Conecte PlayerScore ao TowerPlacement e ao TowerUpgrade.',
    'Configure os tiles de posicionamento válidos com a layer correta.',
    'Ao clicar em uma célula válida com recursos suficientes, a torre é instanciada.',
])

# 3.3 Recursos e Derrota
add_heading(doc, '3.3  Recursos e Derrota — PlayerScore e PlayerHealth', 2, C_TD)
add_subheading(doc, 'PlayerScore', C_TD)
add_body(doc, 'Singleton que controla os recursos do jogador (moeda/pontos) usados para comprar e fazer upgrade de torres.')
add_properties_table(doc, C_TD, [
    ('Start Score',      'int | 100',         'Recursos iniciais do jogador'),
    ('Score UI',         'TMP_Text | null',   'Texto que exibe o saldo de recursos'),
    ('On Score Changed', 'UnityEvent<int>',   'Evento disparado ao alterar o saldo'),
])
add_subheading(doc, 'PlayerHealth (Tower Defense)', C_TD)
add_properties_table(doc, C_TD, [
    ('Max Lives',       'int | 20',           'Número de vidas (inimigos que podem passar)'),
    ('Lives UI',        'TMP_Text | null',    'Texto exibindo as vidas restantes'),
    ('Game Over Panel', 'GameObject | null',  'Painel exibido ao perder todas as vidas'),
    ('On Lives Changed','UnityEvent<int>',    'Evento disparado ao perder uma vida'),
])
add_subheading(doc, 'Passo a Passo — Integrando PlayerHealth com EnemyPath', C_TD)
add_steps(doc, [
    'Selecione o prefab do inimigo e encontre o componente EnemyPath.',
    'No campo Reached End (UnityEvent), clique em +.',
    'Arraste o GameObject com PlayerHealth para o campo de objeto.',
    'Selecione PlayerHealth → TakeDamage na lista de funções.',
    'Ao chegar ao fim, o inimigo automaticamente reduz as vidas e se destrói.',
])

# 3.4 TowerShop
add_heading(doc, '3.4  Loja de Torres — TowerShop', 2, C_TD)
add_body(doc, 'O TowerShop gerencia a interface de seleção de torres, permitindo ao jogador escolher qual tipo de torre deseja comprar antes de posicioná-la.')
add_properties_table(doc, C_TD, [
    ('Tower Options',    'TowerData[]',        'Array de torres disponíveis para compra (prefab, custo, ícone)'),
    ('Selected Tower',   'int | -1',           'Índice da torre selecionada (-1 = nenhuma)'),
    ('Button Container', 'Transform | null',   'Pai dos botões gerados dinamicamente para cada torre'),
    ('Player Score',     'PlayerScore | null', 'Referência para verificar saldo antes de liberar a seleção'),
])
add_subheading(doc, 'Passo a Passo', C_TD)
add_steps(doc, [
    'Adicione TowerShop a um GameObject na cena.',
    'Preencha o array Tower Options com os ScriptableObjects de cada torre disponível.',
    'Atribua Button Container a um HorizontalLayoutGroup ou VerticalLayoutGroup no Canvas.',
    'Conecte Player Score para que a loja verifique o saldo antes de permitir a seleção.',
])
add_tip(doc, 'Fluxo completo: EnemySpawner → EnemyPath → PlayerHealth (derrota) | PlayerScore → TowerShop → TowerPlacement → TowerShooter → TowerUpgrade.')

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# PARTE 4 — POINT & CLICK
# ═══════════════════════════════════════════════════════════════════════════════
add_part_banner(doc, 4, 'Point & Click', C_PAC)
add_body(doc,
    'O módulo Point & Click fornece os fundamentos para jogos de aventura e educativos: '
    'inventário com arrastar e soltar, combinação de itens, grafo de estados de cena, '
    'itens em queda, atirador de precisão e sistemas de fila de comandos para mecânicas '
    'de programação educacional.')

# 4.1 Itens
add_heading(doc, '4.1  Itens — Item e CombinationRecipe', 2, C_PAC)
add_body(doc, 'Os itens são ScriptableObjects criados em Assets → Create → Thinklib → Point and Click → Item.')
add_properties_table(doc, C_PAC, [
    ('Item Name',      'string',        'Nome do item exibido no inventário'),
    ('Icon',           'Sprite | null', 'Sprite exibido no slot do inventário'),
    ('Description',    'string',        'Texto descritivo do item'),
    ('Is Combinable',  'bool | false',  'Se verdadeiro, pode participar de receitas de combinação'),
    ('Is Stackable',   'bool | false',  'Se verdadeiro, itens do mesmo tipo se acumulam em um slot'),
    ('Has Timer',      'bool | false',  'Se verdadeiro, o item tem vida útil limitada na Dropzone'),
    ('Item Lifetime',  'float | 10',   'Duração em segundos antes do item expirar na Dropzone'),
])
add_subheading(doc, 'CombinationRecipe', C_PAC)
add_body(doc, 'As receitas são ScriptableObjects criados em Assets → Create → Thinklib → Point and Click → Combination Recipe.')
add_properties_table(doc, C_PAC, [
    ('Ingredient A',    'Item | null',  'Primeiro item necessário para a combinação'),
    ('Ingredient B',    'Item | null',  'Segundo item necessário'),
    ('Result Item',     'Item | null',  'Item resultante da combinação'),
    ('Consume Inputs',  'bool | true',  'Remove os ingredientes do inventário após combinar'),
])
add_subheading(doc, 'Passo a Passo', C_PAC)
add_steps(doc, [
    'Em Assets, clique com o botão direito → Create → Thinklib → Point and Click → Item.',
    'Preencha Item Name, Icon e Description.',
    'Para itens combináveis, marque Is Combinable = true.',
    'Crie uma CombinationRecipe e atribua Ingredient A, Ingredient B e Result Item.',
    'Os ScriptableObjects são compartilhados entre objetos e cenas, facilitando o reuso.',
])
add_tip(doc, 'Crie os ScriptableObjects de Item e CombinationRecipe antes de montar a cena, pois são referenciados pelo InventoryManager e pelas Dropzones.')

# 4.2 Inventário e Dropzone
add_heading(doc, '4.2  Inventário e Dropzone', 2, C_PAC)
add_body(doc, 'O sistema de inventário usa slots (ItemSlot) onde o jogador arrasta e solta itens. A Dropzone recebe itens do inventário e valida soluções de puzzle.')
add_subheading(doc, 'InventoryManager', C_PAC)
add_properties_table(doc, C_PAC, [
    ('Slot Count',       'int | 8',              'Número de slots do inventário'),
    ('Slot Prefab',      'GameObject | null',    'Prefab de cada slot visual'),
    ('Recipes',          'CombinationRecipe[]',  'Lista de receitas ativas na cena'),
    ('On Item Combined', 'UnityEvent<Item>',     'Evento disparado ao combinar itens com sucesso'),
])
add_subheading(doc, 'DropZone', C_PAC)
add_properties_table(doc, C_PAC, [
    ('Zone ID',        'int',                  'ID único e ordenado da zona (0 para a primeira)'),
    ('Display Sprite', 'SpriteRenderer | null','SpriteRenderer filho usado para exibir o item colocado'),
])
add_subheading(doc, 'ItemSlot (melhorias v0.2.1)', C_PAC)
add_body(doc, 'O ItemSlot exibe a quantidade de itens empilháveis e um contador de timer em vermelho para itens com Has Timer = true.')
add_properties_table(doc, C_PAC, [
    ('Icon Image',       'Image | null',      'Imagem do ícone do item'),
    ('Slot Background',  'Image | null',      'Fundo do slot (muda de cor quando selecionado)'),
    ('Value Text',       'TMP_Text | null',   'Exibe o valor do item ou score do pawn ativo'),
    ('Quantity Text',    'TMP_Text | null',   'Exibe quantidade (stackable) ou timer restante'),
    ('Timer Color',      'Color | vermelho',  'Cor do texto de timer para itens com hasTimer'),
])
add_subheading(doc, 'Passo a Passo', C_PAC)
add_steps(doc, [
    'Adicione InventoryManager a um GameObject na cena e configure Slot Count e Slot Prefab.',
    'No Canvas, crie um GridLayoutGroup para conter os ItemSlots.',
    'Para cada Dropzone, crie um GameObject com Collider2D e adicione DropZone.',
    'Atribua um Display Sprite (SpriteRenderer filho) a cada DropZone.',
    'Configure DropZoneManager com o array de zonas em ordem para validar o puzzle.',
    'Para itens com timer: marque Has Timer = true no Item e defina Item Lifetime.',
])
add_tip(doc, 'Ao completar o puzzle, o DropZoneManager para automaticamente todos os timers ativos nas zonas.')

# 4.3 Grafo de Cena
add_heading(doc, '4.3  Grafo de Cena — Graph', 2, C_PAC)
add_body(doc, 'O módulo Graph permite criar uma estrutura de estados/cenas interconectadas (típica de jogos point & click), onde cada nó representa um local e as arestas representam as transições possíveis.')
add_properties_table(doc, C_PAC, [
    ('Nodes',           'GraphNode[]',       'Lista de nós do grafo (cada um com nome, sprite de fundo e saídas)'),
    ('Start Node',      'GraphNode | null',  'Nó inicial ao carregar a cena'),
    ('Transition Anim', 'bool | true',       'Exibe animação de fade ao transitar entre nós'),
])
add_subheading(doc, 'Passo a Passo', C_PAC)
add_steps(doc, [
    'Crie ScriptableObjects GraphNode em Assets → Create → Thinklib → Point and Click → Graph Node.',
    'Defina o nome, sprite de fundo e as saídas possíveis de cada nó.',
    'Adicione o componente Graph a um GameObject na cena.',
    'Atribua todos os nós ao array Nodes e defina o Start Node.',
    'As transições ocorrem ao clicar nos objetos interativos configurados nos nós.',
])

# 4.4 FallingItem
add_heading(doc, '4.4  Itens em Queda — FallingItem, ItemSpawner e ItemCollector', 2, C_PAC)
add_body(doc, 'Sistema para minigames de coleta: itens caem de cima e o jogador os coleta movendo um coletor com o mouse. Os itens coletados são adicionados ao InventoryManager.')
add_subheading(doc, 'FallingItem', C_PAC)
add_properties_table(doc, C_PAC, [
    ('Fall Speed',           'float | 5.0',           'Velocidade de queda do item em unidades/s'),
    ('Item Sprite Renderer', 'SpriteRenderer | null', 'Renderer do sprite do item (buscado automaticamente em filhos se não atribuído)'),
])
add_subheading(doc, 'ItemSpawner', C_PAC)
add_properties_table(doc, C_PAC, [
    ('Falling Item Prefab', 'GameObject | null', 'Prefab do objeto que cai (deve ter FallingItem)'),
    ('Items To Spawn',      'List<Item>',        'Lista de itens possíveis a serem spawnados aleatoriamente'),
    ('Spawn Interval',      'float | 2.0',       'Intervalo em segundos entre cada spawn'),
    ('Spawn Area Width',    'float | 10',        'Largura da área de spawn horizontal'),
    ('Initial Delay',       'float | 1.0',       'Atraso inicial antes de começar a spawnar'),
])
add_subheading(doc, 'ItemCollector', C_PAC)
add_properties_table(doc, C_PAC, [
    ('Follow Mouse', 'bool | true',  'Se verdadeiro, o coletor segue a posição horizontal do mouse'),
    ('Min X',        'float | -8',   'Limite esquerdo do movimento horizontal do coletor'),
    ('Max X',        'float | 8',    'Limite direito do movimento horizontal do coletor'),
])
add_subheading(doc, 'Passo a Passo', C_PAC)
add_steps(doc, [
    'Crie um prefab FallingItemPrefab com SpriteRenderer (filho), Rigidbody2D (kinematic) e Collider2D (trigger). Adicione FallingItem.',
    'Crie um GameObject ItemSpawner no topo da tela e configure Falling Item Prefab e Items To Spawn.',
    'Crie um GameObject Collector na parte inferior/média da tela com Collider2D (trigger) e adicione ItemCollector.',
    'Crie um GameObject vazio com tag "DespawnZone" e Collider2D (trigger) abaixo da tela para destruir itens não coletados.',
    'Os itens coletados são automaticamente adicionados ao InventoryManager.instance.',
])
add_tip(doc, 'O ItemSpawner usa InvokeRepeating, iniciando após Initial Delay e repetindo a cada Spawn Interval. O Gizmo mostra a área de spawn no Scene View.')

# 4.5 ProjectShooter
add_heading(doc, '4.5  Atirador — PlayerShooter, Projectile, Target e ScoreManager', 2, C_PAC)
add_body(doc, 'Sistema de atirador simples para minigames: o jogador mira com o mouse e dispara projéteis que acertam alvos, somando pontos ao ScoreManager.')
add_subheading(doc, 'PlayerShooter', C_PAC)
add_properties_table(doc, C_PAC, [
    ('Projectile Prefab', 'GameObject | null', 'Prefab do projétil a ser disparado'),
    ('Fire Point',        'Transform | null',  'Ponto de origem do projétil'),
])
add_subheading(doc, 'Projectile', C_PAC)
add_properties_table(doc, C_PAC, [
    ('Speed',    'float | 20',  'Velocidade de deslocamento do projétil'),
    ('Lifetime', 'float | 5.0','Tempo em segundos antes do projétil se auto-destruir'),
])
add_subheading(doc, 'Target', C_PAC)
add_properties_table(doc, C_PAC, [
    ('Score Value', 'int | 10', 'Pontuação adicionada ao ScoreManager ao acertar este alvo'),
])
add_subheading(doc, 'ScoreManager', C_PAC)
add_properties_table(doc, C_PAC, [
    ('Score Text', 'TMP_Text | null', 'Texto da UI que exibe a pontuação atual'),
])
add_subheading(doc, 'Passo a Passo', C_PAC)
add_steps(doc, [
    'Crie um prefab Projectile com Rigidbody2D, Collider2D (trigger) e o script Projectile.',
    'No GameObject do atirador, adicione PlayerShooter. Crie um filho vazio FirePoint na ponta da arma.',
    'Configure Projectile Prefab e Fire Point no Inspector do PlayerShooter.',
    'Para cada alvo, crie um GameObject com Collider2D (trigger) e adicione Target. Defina Score Value.',
    'Adicione ScoreManager a um GameObject na cena e conecte Score Text ao TMP_Text da UI.',
    'Teste: o atirador rotaciona em direção ao mouse e dispara ao pressionar Fire1 (botão esquerdo).',
])
add_tip(doc, 'O Projectile se destrói ao tocar qualquer objeto. Use layers e LayerMask no Collider do projétil para evitar que colida com o próprio jogador.')

# 4.6 CommandsQueue
add_heading(doc, '4.6  Fila de Comandos — CommandsQueue', 2, C_PAC)
add_body(doc, 'Sistema de programação educacional estilo "Hora do Código": o jogador monta uma sequência de comandos (Mover, Virar Esquerda, Virar Direita) e os executa todos de uma vez. Adicionado na v0.2.1.')
add_subheading(doc, 'PlayerAgent', C_PAC)
add_properties_table(doc, C_PAC, [
    ('Move Speed',  'float | 3.0',   'Velocidade de deslocamento em cada passo (unidades/s)'),
    ('Turn Speed',  'float | 180.0', 'Velocidade de rotação em graus/s'),
])
add_subheading(doc, 'CommandQueueManager', C_PAC)
add_properties_table(doc, C_PAC, [
    ('Player Agent',       'PlayerAgent | null',    'Referência ao agente que executa os comandos'),
    ('Command List Text',  'TMP_Text | null',       'Texto da UI que lista os comandos adicionados'),
])
add_body(doc, 'Comandos disponíveis via botões de UI: AddMoveCommand(), AddTurnLeftCommand(), AddTurnRightCommand(). Executar com RunQueue(), resetar com ClearQueue().')
add_subheading(doc, 'Passo a Passo', C_PAC)
add_steps(doc, [
    'Crie um GameObject Player com SpriteRenderer. Adicione PlayerAgent.',
    'Crie um GameObject CommandManager na cena e adicione CommandQueueManager.',
    'Conecte o PlayerAgent ao campo Player Agent do CommandQueueManager.',
    'No Canvas, crie botões de UI e conecte-os: OnClick → CommandQueueManager.AddMoveCommand(), AddTurnLeftCommand(), AddTurnRightCommand().',
    'Crie um botão "Executar" e conecte a RunQueue(), e um botão "Limpar" conectado a ClearQueue().',
    'Crie um TMP_Text para Command List Text para visualizar os comandos na fila.',
])
add_tip(doc, 'Os comandos já executados aparecem em cinza na lista de UI. ClearQueue() reseta o agente à posição e rotação originais.')

# 4.7 GridCoordinates
add_heading(doc, '4.7  Grid de Coordenadas — GridCoordinates', 2, C_PAC)
add_body(doc, 'Sistema de navegação por coordenadas (linha, coluna) em um grid 2D. O jogador define o destino e o agente se move célula a célula até alcançá-lo. Adicionado na v0.2.1.')
add_subheading(doc, 'GridManager', C_PAC)
add_body(doc, 'Singleton que converte coordenadas (linha, coluna) em posição de mundo.')
add_properties_table(doc, C_PAC, [
    ('Grid Origin', 'Vector2',    'Posição (X, Y) da célula (0, 0) no canto superior esquerdo do grid'),
    ('Cell Size',   'float | 1', 'Tamanho de cada célula do grid em unidades do Unity'),
])
add_subheading(doc, 'GridAgent', C_PAC)
add_properties_table(doc, C_PAC, [
    ('Move Speed', 'float | 5.0', 'Velocidade de deslocamento entre células'),
])
add_subheading(doc, 'GridCommandManager', C_PAC)
add_properties_table(doc, C_PAC, [
    ('Grid Agent',        'GridAgent | null',  'Referência ao agente de grid'),
    ('Command List Text', 'TMP_Text | null',   'Texto que lista os comandos de coordenadas'),
    ('Row Input',         'TMP_InputField',    'Campo de input para a linha de destino'),
    ('Col Input',         'TMP_InputField',    'Campo de input para a coluna de destino'),
])
add_subheading(doc, 'Passo a Passo', C_PAC)
add_steps(doc, [
    'Crie um GameObject GridManager na cena, adicione o script GridManager e configure Grid Origin e Cell Size.',
    'Crie um GameObject Player com SpriteRenderer e adicione GridAgent.',
    'Crie um GameObject GridManager_UI e adicione GridCommandManager.',
    'Conecte Grid Agent ao GridCommandManager.',
    'No Canvas, crie dois TMP_InputField (para Linha e Coluna) e um botão "Adicionar Comando" conectado a AddCommandFromUI().',
    'Crie botões "Executar" (RunQueue) e "Limpar" (ClearQueue).',
    'O agente se move primeiro ajustando todas as linhas, depois todas as colunas.',
])
add_tip(doc, 'GridManager.GetWorldPosition(row, col) pode ser chamado de qualquer script para obter a posição de mundo de uma célula específica do grid.')

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# TELEMETRIA
# ═══════════════════════════════════════════════════════════════════════════════
add_part_banner(doc, 5, 'Telemetria — Sistema MIA', C_TELEMETRY)
add_body(doc,
    'O MIA (Métricas e Instrumentação Analítica) é o sistema de telemetria unificado introduzido '
    'na v0.2.0. Todos os scripts listados nas partes anteriores já enviam eventos automaticamente '
    'ao backend configurado.')

add_subheading(doc, 'Configuração via ThinklibTelemetryConfig', C_TELEMETRY)
add_body(doc, 'Crie o asset de configuração em Assets → Create → Thinklib → Telemetry Config. Preencha os campos:')
add_properties_table(doc, C_TELEMETRY, [
    ('Api Base',      'string',         'URL base do servidor de telemetria'),
    ('Route',         'string | /events','Rota do endpoint que recebe os eventos'),
    ('Batch Window',  'float | 5.0',    'Janela de tempo em segundos para envio em lote'),
    ('Enable Logs',   'bool | true',    'Exibe logs detalhados no Console (payload, status, latência)'),
])

add_subheading(doc, 'Eventos Padrão', C_TELEMETRY)
add_properties_table(doc, C_TELEMETRY, [
    ('mechanic_instantiated', 'Automático', 'Disparado ao instanciar qualquer mecânica Thinklib'),
    ('mechanic_used',         'Automático', 'Disparado ao usar uma mecânica pela primeira vez na sessão'),
    ('mechanic_error',        'Automático', 'Disparado ao ocorrer exceção em uma mecânica instrumentada'),
])

add_subheading(doc, 'Scripts instrumentados por módulo', C_TELEMETRY)
add_bullets(doc, [
    'Common / Effects: DeathEffect, PlayerHurtEffect',
    'Common / LifeSystem: LifeUIBar, LifeUIIcons, UILockerAndFollower',
    'Common / Enviroment: RewardChest',
    'Platformer / Collectibles: CollectibleItem, GameManager',
    'Platformer / Combat: PlatformerProjectileAttackController, PlayerMeleeAttackController, PlayerShooterController, ProjectileDamageDealer',
    'Platformer / Enemy / Patroller: DamageOnTouch, PatrollerAI',
    'Platformer / Enemy / Shooter: EnemyShooterAI',
    'Platformer / Environment: MovingPlatform, TimedPlatform',
    'Platformer / Enviroment: SawHazard',
    'Platformer / Movement: PlatformerJumpController, PlatformerMovementController',
    'Topdown / Combat: PlayerTopdownMeleeAttackController, PlayerTopdownShooterController',
    'Topdown / Enemy / Patroller: TopdownDamageOnTouch, TopdownPatrollerAI',
    'Topdown / Enemy / Shooter: TopdownEnemyShooterAI',
    'Topdown / Movement: TopdownMovementController',
    'Topdown / NPC: DialogueBubble, TopdownNPCController',
    'TowerDefense: PlayerHealth, Bullet, TowerPlacement, TowerShooter, EnemyHealth, EnemyPath, EnemySpawner, PlayerScore, TowerShop, TowerUpgrade',
    'PointAndClick / CommandsQueue: PlayerAgent, CommandQueueManager',
    'PointAndClick / GridCoordinates: GridAgent, GridCommandManager, GridManager',
    'PointAndClick / Dropzone: DropZone, ItemSlot',
])
add_tip(doc, 'O campo plataform (com essa grafia) é o identificador de plataforma enviado ao backend conforme o contrato da API. Não altere o nome do campo ao implementar o servidor receptor.')
add_tip(doc, 'Para desabilitar a telemetria em builds de produção, basta não incluir o ThinklibTelemetryConfig.asset na build ou deixar o campo Api Base vazio.')

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# APÊNDICE — CHECKLISTS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'Apêndice — Checklist de Configuração', 1, C_PLATFORMER)

add_subheading(doc, 'Platformer — Checklist', C_PLATFORMER)
add_bullets(doc, [
    'Rigidbody2D com Gravity Scale > 0 e Freeze Rotation Z',
    'CapsuleCollider2D ou BoxCollider2D no personagem',
    'PlatformerMovementController + PlatformerJumpController',
    'GroundCheck filho posicionado na base do collider',
    'Ground Layer configurada nos tiles/plataformas',
    'Animator com parâmetros Speed (float) e IsGrounded (bool)',
    'Sistema de vida: LifeUIBar ou LifeUIIcons no Canvas',
    'FirePoint filho para ataques à distância',
    'Pontos de patrulha configurados nos inimigos PatrollerAI',
    'Tag "Player" no personagem para detecção por RewardChest e SawHazard',
    'LifeSystemController no personagem para integração com RewardChest e SawHazard',
])

add_subheading(doc, 'Topdown — Checklist', C_TOPDOWN)
add_bullets(doc, [
    'Rigidbody2D com Gravity Scale = 0',
    'TopdownMovementController com Blend Params configurados',
    'Animator com Blend Tree 2D (MoveX, MoveY)',
    'Canvas World Space para DialogueBubble dos NPCs',
    'MainCamera com tag "MainCamera" para mira por mouse',
    'Layers separadas para Player, Enemy e NPC',
])

add_subheading(doc, 'Tower Defense — Checklist', C_TD)
add_bullets(doc, [
    'PlayerScore Singleton na cena (único)',
    'PlayerHealth conectado ao evento Reached End do EnemyPath',
    'Waypoints do EnemyPath configurados no prefab do inimigo',
    'Layer de posicionamento válido configurada nos tiles',
    'TowerShop com lista de TowerData preenchida',
    'EnemySpawner com prefab, spawn point e configuração de ondas',
    'Game Over Panel e Lives UI conectados ao PlayerHealth',
])

add_subheading(doc, 'Point & Click — Checklist', C_PAC)
add_bullets(doc, [
    'ScriptableObjects de Item criados e configurados',
    'CombinationRecipes com ingredientes e resultado definidos',
    'InventoryManager com Slot Count e Slot Prefab configurados',
    'Recipes array no InventoryManager com todas as receitas ativas',
    'DropZoneManager com zonas ordenadas para validação do puzzle',
    'Tag "DespawnZone" configurada para a área de saída (FallingItem)',
    'GridManager na cena (singleton) configurado antes de usar GridAgent',
    'PlayerAgent e CommandQueueManager conectados para uso de CommandsQueue',
    'ScoreManager na cena para uso do sistema de atirador (ProjectShooter)',
])

# ─── SALVAR ───────────────────────────────────────────────────────────────────
output_path = r'c:\Lib\Dev_thinklib\thinklib-core\Thinklib_Core_Manual_v0.2.1.docx'
doc.save(output_path)
print(f'Manual gerado: {output_path}')
